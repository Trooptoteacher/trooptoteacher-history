# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1B,0x2A,0x4A); RED=RGBColor(0xB2,0x22,0x34); GOLD=RGBColor(0xC8,0x9B,0x3C)
INK=RGBColor(0x22,0x22,0x22); MUTED=RGBColor(0x5A,0x65,0x79)
CARD="F7F5EF"

SP='/tmp/claude-0/-home-user-trooptoteacher-history/56efdd95-cf8e-5821-8c3d-6afe28c88c95/scratchpad/'
quiz=json.load(open(SP+'quiz_with_exp.json'))
exits=json.load(open(SP+'exit_tickets.json'))
# US.45 quiz key (kept as reference in student book)
quiz['US.45']=[{'dok':2,'id':'q-us45-dok2-1-b','ans':'A','stem':'Treaty of Versailles and the rise of Hitler','explanation':'Harsh reparations and humiliation fueled resentment that the Nazis exploited.'},
 {'dok':2,'id':'q-us45-dok2-3-b','ans':'B','stem':"Japan's militarism in the 1930s",'explanation':'Military leaders seized control and pursued aggressive expansion.'},
 {'dok':3,'id':'u1-us45-dok3-tc','ans':'A','stem':'Economic crisis and the rise of dictators','explanation':'Depression, post-WWI disappointment, and political chaos let dictators promise order.'},
 {'dok':2,'id':'q-us45-dok2-4','ans':'A','stem':'Fascism vs. communism','explanation':'Fascism promoted extreme nationalism and a corporate state; communism sought a classless society.'}]

TITLES={'US.45':'Rise and Spread of Fascism, Communism, and Totalitarianism','US.46':"FDR's Response to World Crises",
'US.47':'U.S. Response to the Holocaust','US.48':'American Entry into World War II','US.49':'WWII Leaders & Key Figures',
'US.50':'Major WWII Battles: Midway, Iwo Jima, Okinawa & D-Day','US.51':'Special Fighting Forces in WWII',
'US.52':'Women in World War II','US.53':'African Americans in WWII','US.54':'Japanese American Internment & Korematsu',
'US.55':'Home Front Mobilization','US.56':'Manhattan Project & Atomic Bomb','US.57':'Yalta and Potsdam Conferences',
'US.58':'Founding of the United Nations & Cordell Hull'}
EXIT_RAT={
'US.45':'Fascism combined ultranationalism, authoritarian leadership, militarism, and scapegoating under a cult of personality.',
'US.46':'The Quarantine Speech urged peace-loving nations to act together against aggressors — a first step away from strict isolationism.',
'US.47':'U.S. troops liberated camps such as Buchenwald and Dachau in 1945, documenting Nazi atrocities for the world.',
'US.48':'Aggression in both theaters (Axis expansion in Europe; Japan’s expansion and Pearl Harbor) pushed the U.S. to war.',
'US.49':'Commanders coordinated combined-arms operations across theaters (e.g., Eisenhower in Europe, MacArthur in the Pacific).',
'US.50':'Intelligence and code-breaking (notably at Midway) let smaller forces anticipate and defeat the enemy.',
'US.51':'The Tuskegee Airmen compiled a strong bomber-escort record, challenging segregation with proven excellence.',
'US.52':'With men overseas and labor demand high, millions of women entered defense industries and the military.',
'US.53':'Wartime defense jobs drew ~1.5 million African Americans to Northern and Western cities (a second Great Migration).',
'US.54':'Wartime fear and racial prejudice, framed as “military necessity,” drove EO 9066 and mass internment.',
'US.55':'Rationing, bond drives, and factory conversion redirected the whole economy toward war production.',
'US.56':'Oak Ridge, TN was chosen for its isolation, TVA electricity, and rail access to enrich uranium in secret.',
'US.57':'Yalta divided Germany, promised free elections in Eastern Europe, and set up the United Nations.',
'US.58':'The League failed partly because the U.S. never joined; the UN gave major powers a veto to keep them engaged.'}
RETEACH={
'US.45':'Build a cause→effect chain (WWI economic collapse → fear → dictators) and a fascism-vs-communism T-chart.',
'US.46':'Timeline the shift from neutrality to involvement: Quarantine (1937) → Cash-and-Carry (1939) → Lend-Lease, Four Freedoms, Atlantic Charter (1941).',
'US.47':'Sequence the U.S. response before / during / after the war; discuss what more could have been done and why it wasn’t.',
'US.48':'Map the road to Pearl Harbor: oil embargo + Axis expansion → attack → declaration; close-read the “Day of Infamy” speech.',
'US.49':'Build a leader matrix (role · theater · key decision) for Eisenhower, MacArthur, Marshall, and Churchill.',
'US.50':'Annotate a Pacific/Europe map: for each battle, label the geographic or strategic factor that decided it.',
'US.51':'Connect each special unit’s service to the Double V idea — proving capability while facing discrimination.',
'US.52':'Before/after chart of women’s roles; distinguish what changed permanently from what reverted post-war.',
'US.53':'Trace the through-line: Double V → EO 8802 → FEPC → 1948 armed-forces integration.',
'US.54':'Read EO 9066 and the Korematsu holding; hold a structured debate on security vs. civil liberties; note the 1988 apology.',
'US.55':'List home-front sacrifices and match each to the specific war need it served.',
'US.56':'Use a pro/con evidence organizer to weigh Truman’s decision; anchor the Tennessee connection at Oak Ridge.',
'US.57':'Compare Yalta vs. Potsdam outcomes side by side and trace the line to Cold War origins.',
'US.58':'Compare the UN and the League of Nations feature-by-feature; highlight Cordell Hull’s role.'}

doc=Document()
# base style
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5); st.font.color.rgb=INK
st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.08
sec=doc.sections[0]
sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.8); sec.left_margin=Inches(0.9); sec.right_margin=Inches(0.9)
sec.header_distance=Inches(0.5); sec.footer_distance=Inches(0.5)

def run(p,text,size=10.5,color=INK,bold=False,italic=False,font='Calibri'):
    r=p.add_run(text); r.font.name=font; r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold; r.italic=italic; return r
def para(space_before=0,space_after=4,align=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(space_before); p.paragraph_format.space_after=Pt(space_after)
    if align is not None: p.alignment=align
    return p
def h1(text):
    p=para(10,6); run(p,text,20,NAVY,bold=True,font='Calibri')
def h2(text,before=12):
    p=para(before,4); run(p,text,14,NAVY,bold=True)
def h3(text,before=8):
    p=para(before,2); run(p,text,11.5,RED,bold=True)
def body(text,**kw):
    p=para(0,4); run(p,text,**kw); return p
def shade(p,color=CARD):
    pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd')
    for k,v in (('w:val','clear'),('w:color','auto'),('w:fill',color)): sh.set(qn(k),v)
    pPr.append(sh)

# header/footer
def set_hf(section):
    h=section.header; h.is_linked_to_previous=False
    hp=h.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run(hp,'U.S. History Hack™ · Unit 6 · Teacher Guide',9,GOLD,bold=True,font='Calibri')
    f=section.footer; f.is_linked_to_previous=False
    fp=f.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run(fp,'U.S. History Hack™ · Unit 6 (Course Standard) · © 2026 TroopToTeacher Technologies LLC   |   Page ',8.5,MUTED)
    fld1=OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'),'PAGE'); r=OxmlElement('w:r'); rpr=OxmlElement('w:rPr'); r.append(rpr); fld1.append(r); fp._p.append(fld1)
set_hf(sec)

def pagebreak(): doc.add_page_default() if False else doc.add_paragraph().add_run().add_break(6)

from docx.enum.text import WD_BREAK
def newpage():
    p=doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

# ---------------- COVER ----------------
para(60)
p=para(0,2,WD_ALIGN_PARAGRAPH.CENTER); run(p,'U.S. HISTORY HACK™',22,NAVY,bold=True)
p=para(0,2,WD_ALIGN_PARAGRAPH.CENTER); run(p,'UNIT 6 · World War II, 1939–1945',14,RED,bold=True)
p=para(10,2,WD_ALIGN_PARAGRAPH.CENTER); run(p,'Teacher How-to-Use & MTSS Guide',18,NAVY,bold=True)
p=para(0,2,WD_ALIGN_PARAGRAPH.CENTER); run(p,'Course Standard Edition · US.45–US.58',12,MUTED)
p=para(16,2,WD_ALIGN_PARAGRAPH.CENTER); run(p,'Answer keys · exit-ticket keys · CER rubric · UDL/MTSS supports · pacing · reteach',11,INK,italic=True)
p=para(40,2,WD_ALIGN_PARAGRAPH.CENTER); run(p,'Companion to the Unit 6 Student Workbook (Course Standard). Teacher edition — keep separate from students.',10.5,MUTED,italic=True)
p=para(30,2,WD_ALIGN_PARAGRAPH.CENTER); run(p,'© 2026 TroopToTeacher Technologies LLC · Aligned to TN Academic Standards & TCAP EOC',9.5,MUTED)

# ---------------- HOW TO USE ----------------
newpage()
h1('How to Use This Workbook')
body('Each standard (US.45–US.58) runs the same seven-activity cycle so students always know what comes next. '
     'The workbook is one common book designed for learner variability: every student follows the CORE PATH, and the '
     'optional supports flex access without lowering the goal.', size=10.5)
h3('The seven-activity cycle')
for t in ['1. Vocabulary (reference / word bank)','2. Vocabulary Studio (Frayer-inspired)','3. Direct-Teaching Cornell Notes',
          '4. Close Read','5. Primary Source / Data Analysis (HIPPO)','6. Core Application: Practice Quiz','7. Constructed Response (CER)']:
    p=para(0,1); run(p,'•  '+t,10.5)
body('Each standard also opens with a warm-up (Set Your Goal, Hook, Activate, Preview & Predict) and closes with an Exit Ticket.', size=10.5)

h2('Pacing — three schedule variants')
body('Each activity carries a time chip. Choose the activities that fit your block; the rest become warm-ups, stations, or homework. '
     'Protect the Exit Ticket in every variant — it is your daily check for understanding.')
for name,desc in [('46-minute regular day','Full cycle across the standard; run one to two activities per day plus the Exit Ticket.'),
                  ('43-minute activities schedule','Trim one optional support or move the Vocabulary Studio to homework; keep Close Read, Quiz/CER, and the Exit Ticket.'),
                  ('41-minute late-start','Prioritize CORE PATH: Cornell Notes or Close Read, one application (Quiz or CER), and the Exit Ticket.')]:
    p=para(2,1); run(p,name+'. ',10.5,NAVY,bold=True); run(p,desc,10.5)

h2('Front / back printing (your choice)')
body('Each activity is built as a front (clean, independent CORE PATH) with UDL/MTSS supports on the back page '
     '(Vocabulary Supports, HIPPO Supports, Writing Supports). Print front-only, back-only, or both — the goal on the front never changes.')

# ---------------- UDL / MTSS ----------------
newpage()
h1('UDL · MTSS Implementation')
body('Supports are available by design and work alongside — never in place of — the accommodations and modifications in a '
     'student’s IEP or 504 plan.', bold=False)
for label,desc in [
 ('CORE PATH','The essential instruction every student receives. Same rigor for everyone.'),
 ('Vocabulary Supports (back of Activity 1)','Word-attack strategy, Spanish cognates, and quick practice. Assign for students who need a vocabulary on-ramp.'),
 ('HIPPO Supports (back of Activity 5)','Guiding questions, a sentence frame, a worked model analysis, and a sourcing warm-up. Assign before the front-page source analysis for scaffolded access.'),
 ('Writing Supports (back of Activity 7)','Sentence stems, a worked CER model, the CER rubric in student language, and an argument word bank. Assign for students building argument writing.'),
 ('Response Choice','Students may show learning by writing, saying/recording, or diagramming — the standard is the thinking, not the format.')]:
    h3(label); body(desc)

# ---------------- CER RUBRIC ----------------
newpage()
h1('Constructed Response (CER) — Scoring Rubric')
body('Score each trait 1–4 for a total of /16. Students self-grade against the same criteria on the front of Activity 7.')
rub=[('Claim','Answers the question in one clear, defensible sentence.'),
     ('Evidence','Uses two specific, accurate facts drawn from the sources / standard.'),
     ('Reasoning','Explains HOW the evidence proves the claim (not just restating it).'),
     ('Conventions','Complete, readable sentences a reader can follow.')]
tbl=doc.add_table(rows=1,cols=3); tbl.style='Table Grid'
hdr=tbl.rows[0].cells
for c,txt in zip(hdr,['Trait','What a strong (3–4) response shows','Score']):
    c.paragraphs[0].clear(); run(c.paragraphs[0],txt,10.5,NAVY,bold=True)
for name,desc in rub:
    cells=tbl.add_row().cells
    run(cells[0].paragraphs[0],name,10.5,INK,bold=True)
    run(cells[1].paragraphs[0],desc,10.5)
    run(cells[2].paragraphs[0],'___ / 4',10.5,MUTED)
body('')
p=para(4,2); run(p,'Scoring bands: ',10.5,NAVY,bold=True); run(p,'14–16 secure · 11–13 developing · 8–10 approaching · below 8 reteach.',10.5)

# ---------------- PER-STANDARD KEYS ----------------
newpage()
h1('Per-Standard Answer Keys & Reteach')
body('Keep this section separate from students. Each standard lists the Practice Quiz key (with rationale), the Exit Ticket key, '
     'and a targeted reteach move. Worked CER and HIPPO models appear on the support (back) pages of the student workbook.')
import re as _re
def clean_rationale(exp):
    if not exp: return ''
    # split into sentences, drop generic distractor-dismissal sentences
    sents=_re.split(r'(?<=[.!?])\s+', exp)
    keep=[s for s in sents if not _re.match(r'^[A-D]\s+is\s+incorrect', s.strip()) and 'describes a different program' not in s]
    out=' '.join(keep[:2]).strip()
    if out and out[-1] not in '.!?': out+='.'
    return out
letters=['A','B','C','D']
first=True
for code in ['US.45','US.46','US.47','US.48','US.49','US.50','US.51','US.52','US.53','US.54','US.55','US.56','US.57','US.58']:
    if not first: newpage()
    first=False
    h2(f'{code} — {TITLES[code]}', before=2)
    h3('Practice Quiz (Activity 6) — answer key')
    for n,it in enumerate(quiz[code],1):
        p=para(1,1)
        run(p,f'{n}. ',10.5,INK,bold=True); run(p,f'{it["ans"]}',11,RED,bold=True)
        run(p,f'   (DOK {it["dok"]})  ',9.5,MUTED)
        rat=clean_rationale(it.get('explanation') or '')
        run(p,rat,10)
    h3('Exit Ticket — answer key')
    ev=exits.get(code,{})
    ans=ev.get('clean',['?'])
    ans=ans[-1] if ans else '?'
    p=para(1,1); run(p,'Correct: ',10.5,INK,bold=True); run(p,ans,11,RED,bold=True)
    run(p,'   '+EXIT_RAT.get(code,''),10)
    h3('Reteach move')
    body(RETEACH.get(code,''))

doc.save(SP+'Unit6_Teacher_Guide_CourseStandard.docx')
print('saved teacher guide')
