#!/usr/bin/env python3
"""Editable Word (.docx) versions of the four Homestead handouts, built natively
with python-docx (reuses the verified data from gen_homestead_reading_hipp).

PDF remains the print-first deliverable per the build standard; these .docx are
the optional editable copies Sean requested.
"""
import html as _html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

import gen_homestead_reading_hipp as G  # data: CHART, FULL_SECTIONS, HIPP, PAGE_GROUPS, CITE, ISO, STAMP, OUT

OUT = G.OUT
STAMP = G.STAMP
NAVY, GOLD, RED, CARD, INK = RGBColor(0x1F, 0x3A, 0x5F), RGBColor(0xC9, 0xA2, 0x27), \
    RGBColor(0xB2, 0x22, 0x34), "F8F5EF", RGBColor(0x20, 0x26, 0x2E)
COPYRIGHT = "© 2026 TroopToTeacher Technologies LLC · U.S. History Hack™. All rights reserved. Licensed for single-classroom reproduction."


def clean(s):
    s = re.sub(r"<[^>]+>", "", s)
    return _html.unescape(s).strip()


def shade(el, hex_):
    """Apply background shading to a cell or paragraph via w:shd."""
    pr = el._tc.get_or_add_tcPr() if hasattr(el, "_tc") else el._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_)
    pr.append(shd)


def bottom_border(par):
    pPr = par._p.get_or_add_pPr()
    pbd = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AEB6C2")
    pbd.append(bot)
    pPr.append(pbd)


def base_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Georgia"
    st.font.size = Pt(11)
    for s in doc.sections:
        s.page_height, s.page_width = Inches(11), Inches(8.5)
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.75)
    return doc


def titleblock(doc, kicker, title, sub):
    p = doc.add_paragraph()
    r = p.add_run(kicker); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GOLD; r.font.name = "Arial"
    shade(p, "1F3A5F"); p.paragraph_format.space_after = Pt(0)
    for run_par in [p]:
        pass
    p2 = doc.add_paragraph(); r2 = p2.add_run(title); r2.bold = True; r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r2.font.name = "Georgia"
    shade(p2, "1F3A5F"); p2.paragraph_format.space_after = Pt(0)
    p3 = doc.add_paragraph(); r3 = p3.add_run(sub); r3.italic = True; r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0xDC, 0xE6, 0xF1)
    shade(p3, "1F3A5F"); p3.paragraph_format.space_after = Pt(8)


def heading(doc, text, size=13, fill="1F3A5F", color=RGBColor(0xFF, 0xFF, 0xFF)):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = color; r.font.name = "Arial"
    shade(p, fill)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    return p


def copyright_block(doc):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(clean(COPYRIGHT)); r.font.size = Pt(8.5); r.font.color.rgb = NAVY; r.font.name = "Arial"
    top = OxmlElement("w:pBdr"); t = OxmlElement("w:top")
    t.set(qn("w:val"), "single"); t.set(qn("w:sz"), "12"); t.set(qn("w:space"), "4"); t.set(qn("w:color"), "C9A227")
    top.append(t); p._p.get_or_add_pPr().append(top)


def note(doc, text):
    p = doc.add_paragraph(); r = p.add_run(clean(text)); r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x43, 0x50, 0x6A)


def set_col_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


# ---------------------------------------------------------------------------
def two_column():
    doc = base_doc()
    titleblock(doc, "U.S. HISTORY HACK™ · STANDARD US.01", "The Homestead Act of 1862",
               "An Act to secure Homesteads to actual Settlers on the Public Domain — excerpt")
    d = doc.add_paragraph(); r = d.add_run("Directions: Read the law’s own words on the left. The right column says the same thing in a shorter sentence.")
    r.bold = True; r.font.name = "Arial"; r.font.size = Pt(11); r.font.color.rgb = NAVY
    d.paragraph_format.space_after = Pt(8)
    chart_map = dict(G.CHART)
    for gi, group in enumerate(G.PAGE_GROUPS):
        if gi > 0:
            doc.add_page_break()
        t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
        hdr = t.rows[0].cells
        for c, txt in zip(hdr, ["The law’s own words (Homestead Act, 1862)", "In shorter words"]):
            shade(c, "1F3A5F"); rr = c.paragraphs[0].add_run(txt); rr.bold = True
            rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rr.font.name = "Arial"; rr.font.size = Pt(10.5)
        for num in group:
            sr = t.add_row().cells; a = sr[0].merge(sr[1])
            shade(a, "C9A227"); rr = a.paragraphs[0].add_run(f"Section {num}"); rr.bold = True
            rr.font.color.rgb = NAVY; rr.font.name = "Arial"; rr.font.size = Pt(12)
            for left, right in chart_map[num]:
                cells = t.add_row().cells
                cells[0].paragraphs[0].add_run(clean(left)).font.size = Pt(10.5)
                rp = cells[1].paragraphs[0].add_run(clean(right)); rp.font.size = Pt(11); rp.font.name = "Arial"
                shade(cells[1], "FBFAF6")
        set_col_widths(t, [Inches(3.6), Inches(3.3)])
    note(doc, G.CLOSING_NOTE)
    note(doc, "Source: " + G.CITE)
    copyright_block(doc)
    doc.save(str(OUT / f"US01_Homestead_Reading_TwoColumn_{STAMP}.docx"))


def full_text():
    doc = base_doc()
    titleblock(doc, "U.S. HISTORY HACK™ · STANDARD US.01", "The Homestead Act of 1862",
               "An Act to secure Homesteads to actual Settlers on the Public Domain — excerpt")
    full_map = dict(G.FULL_SECTIONS)
    for gi, group in enumerate(G.PAGE_GROUPS):
        if gi > 0:
            doc.add_page_break()
        for num in group:
            p = doc.add_paragraph()
            lab = p.add_run(f"Section {num}. "); lab.bold = True; lab.font.name = "Arial"; lab.font.color.rgb = NAVY
            lab.font.size = Pt(13)
            body = p.add_run(clean(full_map[num])); body.font.size = Pt(13)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12); p.paragraph_format.line_spacing = 1.4
    note(doc, G.CLOSING_NOTE)
    note(doc, "Source: " + G.CITE)
    copyright_block(doc)
    doc.save(str(OUT / f"US01_Homestead_FullText_{STAMP}.docx"))


def teacher():
    doc = base_doc()
    titleblock(doc, "U.S. HISTORY HACK™ · STANDARD US.01 · TEACHER REFERENCE",
               "The Homestead Act — in Plain Words", "Chunked plain-language version (≈ reading grade 3–4)")
    b = doc.add_paragraph(); shade(b, "B22234"); b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = b.add_run("TEACHER PLANNING REFERENCE — NOT FOR STUDENT DISTRIBUTION")
    rb.bold = True; rb.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rb.font.name = "Arial"; rb.font.size = Pt(11)
    lp = doc.add_paragraph(); lr = lp.add_run("Use this to plan your questions and check understanding. Students read the primary source (the reading and/or the two-column chart) — this simplified version stays with you.")
    lr.font.name = "Arial"; lr.font.size = Pt(10.5); lr.font.color.rgb = NAVY
    lp.paragraph_format.space_before = Pt(6); lp.paragraph_format.space_after = Pt(8)
    for num, chunks in G.CHART:
        h = heading(doc, f"Section {num}", size=12)
        for _left, right in chunks:
            li = doc.add_paragraph(style="List Bullet"); rr = li.add_run(clean(right)); rr.font.size = Pt(12.5)
            li.paragraph_format.space_after = Pt(3)
    note(doc, "This is a plain-language paraphrase for teacher planning — it is not the primary source. The verbatim law is on the student reading.")
    note(doc, "Paraphrase of the Homestead Act of 1862 (12 Stat. 392, public domain), by TroopToTeacher Technologies LLC.")
    copyright_block(doc)
    doc.save(str(OUT / f"US01_Homestead_TEACHER_PlainWords_{STAMP}.docx"))


def hipp():
    doc = base_doc()
    titleblock(doc, "U.S. HISTORY HACK™ · STANDARD US.01", "HIPP Source Analysis", "The Homestead Act of 1862")
    lead = doc.add_paragraph()
    r1 = lead.add_run("HIPP "); r1.bold = True; r1.font.color.rgb = RED; r1.font.name = "Arial"
    r2 = lead.add_run("helps you analyze a primary source. Look back at the reading and, for each part, write what you notice. ")
    r2.font.name = "Arial"; r2.font.size = Pt(11); r2.font.color.rgb = NAVY
    r3 = lead.add_run("We are analyzing the source — not writing an essay yet."); r3.bold = True
    r3.font.name = "Arial"; r3.font.size = Pt(11); r3.font.color.rgb = RED
    nb = doc.add_paragraph(); nr = nb.add_run("Name: ______________________     Class / Period: ______________     Date: ____________")
    nr.font.name = "Arial"; nr.font.size = Pt(10); nr.font.color.rgb = RGBColor(0x43, 0x50, 0x6A)
    for L, name, q, stem in G.HIPP:
        heading(doc, f"{L} — {name}", size=13)
        pq = doc.add_paragraph(); tr = pq.add_run("Think about: "); tr.bold = True; tr.font.name = "Arial"; tr.font.color.rgb = NAVY
        pq.add_run(clean(q)).font.size = Pt(11); pq.paragraph_format.space_after = Pt(2)
        ps = doc.add_paragraph(); shade(ps, "F8F5EF")
        sr = ps.add_run("Sentence starter: "); sr.bold = True; sr.font.name = "Arial"; sr.font.color.rgb = RED
        ps.add_run(clean(stem)).font.name = "Arial"; ps.runs[-1].font.size = Pt(10)
        for _ in range(2):
            rl = doc.add_paragraph(); rl.add_run(" "); bottom_border(rl); rl.paragraph_format.space_after = Pt(6)
    tp = doc.add_paragraph(); shade(tp, "F8F5EF")
    tr = tp.add_run("Tip: use the Act’s numbered Sections as your evidence — point to a Section (for example, “Section 5 shows …”) instead of copying long sentences.")
    tr.font.name = "Arial"; tr.font.size = Pt(9.5); tr.font.color.rgb = NAVY
    copyright_block(doc)
    doc.save(str(OUT / f"US01_Homestead_HIPP_Chart_{STAMP}.docx"))


def main():
    two_column(); full_text(); teacher(); hipp()
    for f in sorted(OUT.glob(f"*_{STAMP}.docx")):
        print("wrote", f.name)


if __name__ == "__main__":
    main()
