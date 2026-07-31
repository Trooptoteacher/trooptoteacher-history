#!/usr/bin/env python3
"""DBQ / Primary-Source Investigation Book — World History Hack, Unit 1 (Age of Revolutions).
Genuine public-domain document excerpts; HIPPO + SOAPS + OPTIC; scaffolded CER; EN/ES companion;
teacher guide + point rubric + Schedule-F self-score (teacher-side). Leak-guarded."""
import json, re, sys, base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# iOS/Apple/Google previewers ignore fixed table layout and autofit columns to their widest glyph
# content, collapsing blank answer cells. A fixed-width transparent PNG is measured by every renderer,
# so it holds each column open. Written once to disk for python-docx's add_picture().
_SPACER_PATH='_dbq_spacer.png'
open(_SPACER_PATH,'wb').write(base64.b64decode(
  'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4nGP4//8/AwAI/AL+p5qgoAAAAABJRU5ErkJggg=='))
def _fixed(t):
    """Force fixed table layout (autofit off) so column widths are honored on desktop viewers."""
    t.autofit=False; t.allow_autofit=False
    tblPr=t._tbl.tblPr
    for el in tblPr.findall(qn('w:tblLayout')): tblPr.remove(el)
    lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
def _bborder(p,color='8C8C8C',sz='8'):
    """Give a paragraph a visible bottom baseline (a writing line)."""
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    for k,v in (('w:val','single'),('w:sz',sz),('w:space','2'),('w:color',color)): b.set(qn(k),v)
    pbdr.append(b); pPr.append(pbdr)
def _writecell(cell,width_in,lines=2):
    """Blank answer cell: invisible width anchor + visible ruled baselines (non-merging via spacers)."""
    p0=cell.paragraphs[0]
    p0.add_run().add_picture(_SPACER_PATH,width=Emu(max(64000,int((width_in-0.2)*914400))),height=Emu(9525))
    for i in range(lines):
        p=p0 if i==0 else cell.add_paragraph()
        _bborder(p)
        if i<lines-1:
            sp=cell.add_paragraph(); sp.paragraph_format.space_before=Pt(0); sp.paragraph_format.space_after=Pt(3)

C=json.load(open('analysis/unit1_content.json')); U=C['unit']
BRAND=U['brand']; NAVY=RGBColor(0x1B,0x2A,0x4A); GOLD=RGBColor(0xC8,0x9B,0x3C); RED=RGBColor(0xB2,0x22,0x34); GREY=RGBColor(0x60,0x60,0x60)
doc=Document()
for s in doc.sections:
    s.top_margin=s.bottom_margin=Inches(0.8); s.left_margin=s.right_margin=Inches(0.85)
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11)

def para(text='',size=11,bold=False,italic=False,color=None,align=None,after=6,before=0):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); return p
def rule(): para('—'*54,size=9,color=GREY,after=4)
def box(rows, widths=(2.4,4.6)):
    t=doc.add_table(rows=0,cols=2); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    _fixed(t)
    for a,b in rows:
        c=t.add_row().cells
        c[0].width=Inches(widths[0]); c[1].width=Inches(widths[1])
        rp=c[0].paragraphs[0].add_run(a); rp.bold=True; rp.font.size=Pt(10)
        if b:
            c[1].paragraphs[0].add_run(b).font.size=Pt(10)
        else:
            _writecell(c[1],widths[1])   # blank answer cell -> width anchor + ruled writing lines
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

# ---- COVER ----
para(BRAND+'™',size=16,bold=True,color=GOLD,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
para('DBQ · Primary-Source Investigation',size=13,bold=True,color=RED,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
para(U['code']+' — '+U['title'],size=26,bold=True,color=NAVY,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
para('Where should the power to rule come from — God, or the people?',size=15,italic=True,color=NAVY,align=WD_ALIGN_PARAGRAPH.CENTER,after=10)
para('Standards W.01 · W.03 · W.08a   ·   Course Standard Edition   ·   Grades 9–12',size=11,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER,after=10)
para('Universal Design for Learning 3.0 (CAST, 2024) + MTSS. Read-aloud on request · key terms glossed (EN/ES) · '
     'answer in writing, speech, or a labeled diagram. Same target for everyone; supports vary the means, not the ceiling.',
     size=10,italic=True,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
para('© 2026 TroopToTeacher Technologies LLC. Documents below are public domain. '
     'Assessment is classroom-formative · pre-field-test.',size=9,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ---- THE INVESTIGATION ----
para('The Historical Question',size=16,bold=True,color=NAVY,after=4)
rule()
para('For centuries, European kings claimed they ruled by the divine right of kings — that God had chosen them, so no '
     'one on earth could question them. During the 1600s and 1700s, Enlightenment thinkers argued something '
     'revolutionary instead: that a government’s power comes from the consent of the governed. In this investigation you '
     'will read three primary sources, analyze each one, and build an evidence-based argument answering:',after=6)
para('Should the power to rule come from God (divine right) or from the people (consent of the governed)?',
    size=13,bold=True,color=RED,after=8)
para('How to work like a historian: for each document, first SOURCE it (who wrote it, when, and why), then read it '
     'closely, then use HIPPO to analyze it. Mark up the text. You will use these documents as evidence in your final argument.',
     italic=True,color=GREY,after=8)

DOCS=[
 {"id":"A","title":"On the Divine Right of Kings","who":"King James I of England","date":"Speech to Parliament, 1610",
  "repo":"Public domain (pre-1929).","tag":"divine right (supports W.01)",
  "text":"“The state of monarchy is the supremest thing upon earth; for kings are not only God’s lieutenants upon earth, "
         "and sit upon God’s throne, but even by God himself they are called gods… Kings are justly called gods, for that "
         "they exercise a manner or resemblance of divine power upon earth.”"},
 {"id":"B","title":"Second Treatise of Government","who":"John Locke","date":"1689",
  "repo":"Public domain (pre-1929).","tag":"natural rights & consent (supports W.03)",
  "text":"“Men being, as has been said, by nature all free, equal, and independent, no one can be put out of this estate "
         "and subjected to the political power of another without his own consent… When any number of men have so consented "
         "to make one community or government… the majority have a right to act and conclude the rest.”"},
 {"id":"C","title":"The Declaration of Independence","who":"Thomas Jefferson and the Continental Congress","date":"1776",
  "repo":"U.S. National Archives. Public domain (US government).","tag":"consent of the governed (supports W.08a)",
  "text":"“We hold these truths to be self-evident, that all men are created equal, that they are endowed by their Creator "
         "with certain unalienable Rights… That to secure these rights, Governments are instituted among Men, deriving their "
         "just powers from the consent of the governed.”"},
]
for d in DOCS:
    doc.add_page_break()
    para(f'Document {d["id"]} — {d["title"]}',size=15,bold=True,color=NAVY,after=2)
    para(f'{d["who"]}, {d["date"]}. {d["repo"]}  ·  Theme: {d["tag"]}',size=10,italic=True,color=GREY,after=6)
    para(d['text'],size=12,after=8)
    para('SOAPS — source it first',size=12,bold=True,color=RED,after=2)
    box([('S — Subject: what is it about?',''),('O — Occasion: what was happening then?',''),
         ('A — Audience: who was meant to hear it?',''),('P — Purpose: why was it written?',''),
         ('S — Speaker: whose point of view is it?','')])
    para('HIPPO — analyze the source',size=12,bold=True,color=RED,after=2)
    box([('H — Historical context',''),('I — Intended audience',''),('P — Purpose',''),
         ('P — Point of view',''),('O — Outside connection (to a unit standard)','')])
    para('Support option — sentence frame: “This source argues ______ because it ______, which reveals ______.”',
         italic=True,color=GREY,size=10,after=2)
    para('Response choice: you may answer these boxes in writing, by recording your spoken answer, or with a labeled diagram.',
         italic=True,color=GREY,size=10)

# OPTIC (for the visual pairing) + comparison
doc.add_page_break()
para('Put the documents in conversation',size=15,bold=True,color=NAVY,after=4); rule()
para('OPTIC (if your teacher pairs these texts with a portrait such as Rigaud’s Louis XIV): Overview · Parts · Title · '
     'Interrelationships · Conclusion — how does the image project the same idea of power that Document A defends?',size=10,italic=True,color=GREY,after=6)
box([('Where A (divine right) and B/C (consent) most directly DISAGREE',''),
     ('Which document’s idea won out in later revolutions — and your evidence',''),
     ('One question you still have',''),], widths=(3.6,3.4))

# ---- CER ARGUMENT ----
doc.add_page_break()
para('Build Your Argument (CER)',size=16,bold=True,color=NAVY,after=4); rule()
para('Should the power to rule come from God (divine right) or from the people (consent of the governed)? '
     'Write a claim, support it with evidence from at least TWO documents, and explain your reasoning.',size=12,bold=True,after=6)
para('Argument frames (use if helpful): “The power to rule should come from ______.”  ·  '
     '“Document ___ supports this because it states ‘______.’”  ·  '
     '“This proves ______ because ______.”  ·  “Although Document ___ argues ______, the stronger case is ______.”',
     italic=True,color=GREY,size=10,after=6)
box([('CLAIM — answer the question',''),('EVIDENCE 1 — quote/cite a document',''),
     ('EVIDENCE 2 — quote/cite a different document',''),('REASONING — how the evidence proves your claim',''),
     ('COUNTERPOINT — address the other side','')], widths=(2.4,4.6))
para('Self-check: Is my claim a clear answer? Do I cite TWO documents by letter? Does my reasoning explain HOW the '
     'evidence proves the claim? Did I address the other side?',italic=True,color=GREY,size=10)

# ---- EN/ES LANGUAGE COMPANION ----
doc.add_page_break()
para('Language Access Companion (EN / ES)',size=15,bold=True,color=NAVY,after=4); rule()
para('Key terms for this investigation, in English and Spanish. Términos clave para esta investigación, en inglés y español.',
    italic=True,color=GREY,size=10,after=6)
vt=doc.add_table(rows=1,cols=3); vt.style='Table Grid'; _fixed(vt)
_VW=(2.0,2.0,2.8)
hdr=vt.rows[0].cells
for i,h in enumerate(['Term / Término','Spanish / Español','Meaning / Significado']):
    hdr[i].width=Inches(_VW[i])
    r=hdr[i].paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=NAVY
GLOSS=[('divine right','derecho divino','the belief a king’s power comes from God'),
       ('absolutism','absolutismo','total, centralized power in one monarch'),
       ('natural rights','derechos naturales','rights all people have by being human (Locke)'),
       ('consent of the governed','consentimiento de los gobernados','government’s power comes from the people'),
       ('unalienable rights','derechos inalienables','rights that cannot be taken away'),
       ('social contract','contrato social','the agreement between rulers and the people')]
for a,b,c in GLOSS:
    cells=vt.add_row().cells
    for i,txt in enumerate([a,b,c]):
        rr=cells[i].paragraphs[0].add_run(txt); rr.font.size=Pt(10)
        if i==0: rr.bold=True
para('')
para('Frame in both languages: “The power to rule should come from ______, because ______.”  /  '
     '“El poder para gobernar debe venir de ______, porque ______.”',italic=True,color=GREY,size=10)

# ---- TEACHER GUIDE (teacher-side) ----
doc.add_page_break()
para('TEACHER GUIDE — keep this section; do not distribute',size=15,bold=True,color=RED,after=4); rule()
para('Purpose & standards. This investigation lets students weigh the divine right of kings (W.01) against the '
     'Enlightenment ideas of natural rights and consent of the governed (W.03) and their spread into the Declaration of '
     'Independence (W.08a). It builds SSP.01–SSP.04 (source, examine, synthesize, argue).',after=6)
para('Document keys / what to look for:',bold=True,after=2)
box([('Doc A (James I, 1610)','Classic divine-right claim: kings answer only to God. Point of view: a reigning monarch defending royal power.'),
     ('Doc B (Locke, 1689)','Consent + majority rule; no legitimate power over a person without consent. The Enlightenment counter-argument.'),
     ('Doc C (Declaration, 1776)','Locke’s ideas in action: unalienable rights + “consent of the governed.” Shows ideas moving into revolution.')],widths=(2.2,5.0))
para('Point rubric (10 pts):',bold=True,after=2)
box([('Claim answers the question (2)',''),('Evidence cites ≥2 documents accurately (3)',''),
     ('Reasoning explains how evidence proves the claim (3)',''),('Addresses the counter-position (1)',''),
     ('Uses academic vocabulary correctly (1)','')],widths=(4.4,2.6))
para('Reteach (MTSS). Tier 2: small-group — re-read Doc B and Doc C side by side, highlight every phrase about "consent," '
     'then rebuild the claim with the frame. Tier 3: 1:1 concrete→representational→abstract — sort the three documents into '
     '"power from God" vs. "power from the people," then co-construct one sentence of reasoning; progress-monitor to the same standard.',after=6)
para('Schedule F self-score (TN textbook adoption):',bold=True,after=2)
box([('Standards alignment — W.01/W.03/W.08a addressed with primary sources','✔'),
     ('Primary sources genuinely public-domain and cited','✔'),
     ('Balanced perspectives (divine right AND consent presented fairly)','✔'),
     ('Accessibility — UDL 3.0 supports + EN/ES companion + response choice','✔'),
     ('Assessment labeled classroom-formative · pre-field-test','✔')],widths=(5.4,1.6))
para('Attribution. Document excerpts are public domain. Any connecting narrative is '+BRAND+'-authored instructional '
     'synthesis, not presented as a primary source.',size=9,italic=True,color=GREY)

out='deliverables/Unit1_DBQ_Investigation_Book.docx'
doc.save(out)
# leak guard on the rendered text
import subprocess, zipfile, os
os.makedirs('deliverables',exist_ok=True)
xml=zipfile.ZipFile(out).read('word/document.xml').decode('utf-8','ignore')
FORB=r'Government Hack|Government & Civics|U\.S\. History\b|Foundations of Constitutional|Williamson|WCS|US\.[0-9]{2}|GC\.[0-9]|flight log|\bcivic\b'
hits=set(re.findall(FORB,xml))
print(f"WROTE {out} · leak {'HITS:'+str(hits) if hits else 'clean'}")
if hits: sys.exit(2)
