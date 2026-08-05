#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Repair workbook .docx masters whose media parts were saved as `*.undefined`
with no matching Default content-type (invalid OPC package — won't open in
python-docx or LibreOffice). Deterministic: sniff each media part's real image
type, rename it, rewrite every .rels target, and register the Default extensions
in [Content_Types].xml. Idempotent. Exit 1 if the output still won't open.

Usage: repair_docx_media.py IN.docx OUT.docx
"""
import sys, os, zipfile, shutil, re

MAGIC = [
    (b'\xff\xd8\xff', 'jpeg', 'image/jpeg'),
    (b'\x89PNG\r\n\x1a\n', 'png', 'image/png'),
    (b'GIF87a', 'gif', 'image/gif'),
    (b'GIF89a', 'gif', 'image/gif'),
    (b'BM', 'bmp', 'image/bmp'),
    (b'\x00\x00\x01\x00', 'ico', 'image/x-icon'),
    (b'<?xml', 'svg', 'image/svg+xml'),
    (b'\x01\x00\x00\x00', 'emf', 'image/x-emf'),
]

def sniff(b):
    for magic, ext, ct in MAGIC:
        if b.startswith(magic):
            return ext, ct
    return None, None

def repair(src, out):
    zin = zipfile.ZipFile(src)
    names = zin.namelist()
    rename = {}      # old part path -> new part path
    exts = {}        # ext -> content-type (Defaults to ensure)
    for n in names:
        if n.startswith("word/media/"):
            base = os.path.basename(n)
            stem, dot, ext = base.partition(".")
            if ext == "" or ext.lower() == "undefined":
                b = zin.read(n)[:16]
                e, ct = sniff(b)
                if not e:
                    print(f"  ! cannot sniff {n} (first bytes {b[:8]!r}) — left as-is")
                    continue
                new = f"word/media/{stem}.{e}"
                rename[n] = new
                exts[e] = ct
            else:
                exts.setdefault(ext.lower(), {
                    "jpeg": "image/jpeg", "jpg": "image/jpeg", "png": "image/png",
                    "gif": "image/gif", "bmp": "image/bmp", "emf": "image/x-emf",
                    "wmf": "image/x-wmf", "svg": "image/svg+xml",
                }.get(ext.lower(), "application/octet-stream"))

    def fix_rels(data):
        for old, new in rename.items():
            # rels targets are relative to /word: "media/X.undefined"
            data = data.replace(("media/" + os.path.basename(old)).encode(),
                                ("media/" + os.path.basename(new)).encode())
        return data

    def fix_content_types(data):
        text = data.decode("utf-8")
        present = set(re.findall(r'Default Extension="([^"]+)"', text))
        inject = ""
        for e, ct in sorted(exts.items()):
            if e not in present:
                inject += f'<Default Extension="{e}" ContentType="{ct}"/>'
        if inject:
            text = text.replace("</Types>", inject + "</Types>", 1)
        return text.encode("utf-8")

    tmp = out + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            name = rename.get(item.filename, item.filename)
            if item.filename == "[Content_Types].xml":
                data = fix_content_types(data)
            elif item.filename.endswith(".rels"):
                data = fix_rels(data)
            zi = zipfile.ZipInfo(name, date_time=item.date_time)
            zi.compress_type = zipfile.ZIP_DEFLATED
            zi.external_attr = item.external_attr
            zout.writestr(zi, data)
    zin.close()
    shutil.move(tmp, out)
    return len(rename), exts

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(__doc__); sys.exit(1)
    src, out = sys.argv[1], sys.argv[2]
    n, exts = repair(src, out)
    print(f"repaired {os.path.basename(src)}: renamed {n} media part(s); defaults ensured: {sorted(exts)}")
    # verify it opens
    try:
        from docx import Document
        d = Document(out)
        print(f"  ✓ python-docx opens it — paragraphs={len(d.paragraphs)} tables={len(d.tables)}")
    except Exception as e:
        print(f"  ✗ still won't open: {e}"); sys.exit(1)
