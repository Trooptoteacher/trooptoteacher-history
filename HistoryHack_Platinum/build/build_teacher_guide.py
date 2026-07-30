# -*- coding: utf-8 -*-
"""Unit 1 Teacher How-to-Use & MTSS Guide — authored to the Unit 5 template.
Holds the SEPARATE answer key (de-biased positions synced to the decks + workbook),
MTSS decision routing, scaffold-fading guidance, and per-standard facilitation refs."""
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import unit1_data as D

OUT='deliverables/Unit1_Teacher_HowToUse_and_MTSS_Guide.docx'
NAVY=RGBColor(0x1A,0x23,0x32); GOLD=RGBColor(0x9A,0x77,0x12); INK=RGBColor(0x1F,0x24,0x30); GREY=RGBColor(0x5C,0x64,0x70)
STDS=['US.01','US.02','US.03','US.04','US.05','US.06','US.07']
# Teacher-deck CFU + Answer-reveal positions (final, de-biased)
TREF={'US.01':(16,17),'US.02':(30,31),'US.03':(46,47),'US.04':(61,62),'US.05':(76,77),'US.06':(91,92),'US.07':(106,107)}

doc=Document(); st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11)
def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)
def para(t,size=11,bold=False,color=INK,align=None,italic=False,after=6,before=0):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); return p

t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('U.S. History Hack™'); r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=GOLD
para('Unit 1 — Teacher How-to-Use & MTSS Guide',20,True,NAVY,WD_ALIGN_PARAGRAPH.CENTER,after=4)
para('The Rise of Industrialization (1877–1900) · Course Standard · US.01–US.07',12,False,INK,WD_ALIGN_PARAGRAPH.CENTER,after=14)
para('Release-Ready · Pilot Edition',12,True,NAVY,WD_ALIGN_PARAGRAPH.CENTER,after=14)
para('TroopToTeacher Technologies LLC · © 2026. Answer key in this guide is for teacher use only — keep separate from student materials.',10,False,GREY,WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

doc.add_heading('How to Use This Unit',level=1)
para('Two synchronized decks plus one Course Standard workbook cover US.01–US.07. The Teacher (Full) Deck (111 slides) is the facilitation deck — it carries Quick Reviews, Confidence Checks, Guided Practice, Word Walls, Checks for Understanding, and Answer Reveals with keys. The Student (Lean) Deck (64 slides) is the projection/print deck — same content, no exposed keys, with an added Key Vocabulary / Language Support slide and a hidden Progress Check per standard. The workbook’s Cornell Universal Fronts cite the exact Student-deck slide ranges.',11)
para('One firm learning goal per standard; flexible, universal means. Supports are available by design and never lower the goal; they work alongside — never in place of — required IEP or 504 accommodations.',11)

doc.add_heading('MTSS — Teacher Decision Band',level=1)
para('Every standard wrap-up shows the same routing cycle:',11,after=4)
para('PROGRESS CHECK → identify barrier → select support → reteach → recheck → extend',12,bold=True,color=NAVY,after=8)
band=[('If the Progress Check shows a gap','Name the barrier (vocabulary, reading load, background, or reasoning). Select a SUPPORT OPTION (Word Wall term, sentence frame, Guided Practice re-run), reteach the specific step, then recheck with the same item type.'),
      ('If students meet the goal','Move to the EXTENSION prompt on the wrap-up (DOK 3 argumentation).'),
      ('Response choice','Accept written, spoken/recorded, or diagrammed responses — all demonstrate the same target.')]
tb=doc.add_table(rows=0,cols=2); tb.style='Table Grid'
for k,v in band:
    c=tb.add_row().cells; shade(c[0],'F1EFE6')
    p=c[0].paragraphs[0]; rr=p.add_run(k); rr.font.bold=True; rr.font.size=Pt(10.5); rr.font.color.rgb=NAVY
    c[1].paragraphs[0].add_run(v).font.size=Pt(10.5)

doc.add_heading('Scaffold Fading',level=1)
for lvl,txt2 in [('Guided','Model + sentence frames + Word Wall visible; use for first exposure or after a missed check.'),
                 ('Light','Frames faded to a checklist; Word Wall available on request.'),
                 ('Independent','Student produces the response with no frame.'),
                 ('Guardrail','Fade Guided → Light → independent as evidence grows; reintroduce support for a new barrier. Required IEP/504 accommodations do not fade unless the authorized team changes them.')]:
    p=para('',11,after=4); r1=p.add_run(lvl+':  '); r1.font.bold=True; r1.font.color.rgb=NAVY; r1.font.size=Pt(11); r2=p.add_run(txt2); r2.font.size=Pt(11)
doc.add_page_break()

# ---------- ANSWER KEY (separate) ----------
doc.add_heading('Answer Key — Checks for Understanding / Progress Checks',level=1)
para('Positions are de-biased and identical across the Teacher CFU/Answer Reveal, the Student Progress Check, and the workbook Transfer Check. Distribution across US.01–US.07: B, D, C, A, D, B, A.',10.5,italic=True,color=GREY,after=8)
ak=doc.add_table(rows=1,cols=5); ak.style='Table Grid'
for j,h in enumerate(['Std','DOK','Key','Teacher slides (CFU / Reveal)','Rationale']):
    shade(ak.rows[0].cells[j],'1A2332'); pp=ak.rows[0].cells[j].paragraphs[0]; rr=pp.add_run(h); rr.font.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for code in STDS:
    item=D.CFU[code]; opts,key=D.debiased_options(code); cfu,rev=TREF[code]
    row=ak.add_row().cells
    row[0].paragraphs[0].add_run(code).font.size=Pt(10)
    row[1].paragraphs[0].add_run('DOK '+str(item['dok'])).font.size=Pt(10)
    rr=row[2].paragraphs[0].add_run(key); rr.font.bold=True; rr.font.size=Pt(11); rr.font.color.rgb=NAVY
    row[3].paragraphs[0].add_run(f'{cfu} / {rev}').font.size=Pt(10)
    row[4].paragraphs[0].add_run(item['why']).font.size=Pt(9.5)
para('Assessment items are pre-field-test (classroom-formative), not a secure TCAP form. Field-test before any secure use (see field-test plan).',9.5,italic=True,color=GREY,before=8)

cp=doc.core_properties
cp.author='TroopToTeacher Technologies LLC'; cp.last_modified_by='TroopToTeacher Technologies LLC'
cp.title='U.S. History Hack — Unit 1 Teacher How-to-Use & MTSS Guide (Pilot)'
cp.created=datetime.datetime(2026,7,25); cp.modified=datetime.datetime(2026,7,25)
doc.save(OUT)
print('SAVED',OUT,'| tables:',len(doc.tables))
