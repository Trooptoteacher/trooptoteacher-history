# -*- coding: utf-8 -*-
"""Unit 1 Course Standard STUDENT WORKBOOK — authored to the Unit 5 platinum template.
No editable source DOCX exists, so this builds one from the verified content:
Cornell Universal Fronts with exact FINAL Student-deck slide refs, 25/75 Cornell notes,
vocabulary, de-biased Transfer Checks (answers kept SEPARATE in the teacher guide),
Representation section, and a real TOC field (built later via UNO). Nothing invented."""
import datetime, docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import unit1_data as D

OUT='deliverables/Unit1_Student_Workbook_CourseStandard.docx'
NAVY=RGBColor(0x1A,0x23,0x32); GOLD=RGBColor(0x9A,0x77,0x12); INK=RGBColor(0x1F,0x24,0x30); GREY=RGBColor(0x5C,0x64,0x70)
STDS=['US.01','US.02','US.03','US.04','US.05','US.06','US.07']

# FINAL Student-deck slide references (read from the built 64-slide deck order)
REF={
 'US.01':{'vocab':3,'dt':'4–7','source':8,'persp':9,'progress':12},
 'US.02':{'vocab':14,'dt':'15–17','source':18,'persp':19,'progress':20},
 'US.03':{'vocab':22,'dt':'23–26','source':27,'persp':28,'progress':29},
 'US.04':{'vocab':31,'dt':'32–35','source':36,'persp':37,'progress':38},
 'US.05':{'vocab':40,'dt':'41–43','source':44,'persp':45,'progress':46},
 'US.06':{'vocab':48,'dt':'49–51','source':52,'persp':53,'progress':54},
 'US.07':{'vocab':56,'dt':'57–60','source':61,'persp':62,'progress':63},
}
# canonical-style assessment IDs (deck-sourced CFU items)
QID={'US.01':'u1-us01-dok3-tc','US.02':'u1-us02-dok2-tc','US.03':'u1-us03-dok3-tc',
     'US.04':'u1-us04-dok2-tc','US.05':'u1-us05-dok2-tc','US.06':'u1-us06-dok2-tc','US.07':'u1-us07-dok2-tc'}

doc=Document()
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11)
for lvl,sz in [('Title',30),('Heading 1',18),('Heading 2',13.5)]:
    s=doc.styles[lvl]; s.font.name='Calibri'
def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)
def para(text,size=11,bold=False,color=INK,align=None,italic=False,after=6,before=0):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); return p

# ---------- COVER ----------
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('U.S. History Hack™'); r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=GOLD
para('Unit 1 — The Rise of Industrialization (1877–1900)',22,True,NAVY,WD_ALIGN_PARAGRAPH.CENTER,after=4)
para('Course Standard · Student Workbook',15,False,INK,WD_ALIGN_PARAGRAPH.CENTER,after=2)
para('Tennessee U.S. History Standards US.01–US.07 · Grades 9–12',12,False,GREY,WD_ALIGN_PARAGRAPH.CENTER,after=18)
para('Release-Ready · Pilot Edition',12,True,NAVY,WD_ALIGN_PARAGRAPH.CENTER,after=18)
para('TroopToTeacher Technologies LLC · © 2026. All rights reserved.',10,False,GREY,WD_ALIGN_PARAGRAPH.CENTER,after=2)
para('Aligned to TN Academic Standards & TCAP EOC. Assessment items are pre-field-test (classroom-formative), not a secure TCAP form.',9.5,False,GREY,WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ---------- TOC placeholder (real field inserted, built via UNO) ----------
doc.add_heading('Table of Contents',level=1)
para('Right-click to update this table of contents (updates on open in Word/LibreOffice).',10,italic=True,color=GREY)
doc.add_page_break()

# ---------- How to use + legend ----------
doc.add_heading('How to Use This Workbook — UDL & MTSS',level=1)
para('One firm learning goal per standard; flexible, universal ways to reach it. These supports are available by design and never lower the goal. They work alongside — never in place of — required IEP or 504 accommodations.',11)
legend=[('CORE PATH','The essential instruction every student completes.'),
        ('SUPPORT OPTION','An optional scaffold that keeps the goal, not lowers it.'),
        ('LANGUAGE SUPPORT','Vocabulary, pronunciation, and Spanish cognates for access.'),
        ('RESPONSE CHOICE','Show learning by writing, saying/recording, or diagramming.'),
        ('PROGRESS CHECK','A quick DOK-2/3 check that guides reteach or extend.'),
        ('EXTENSION','A deeper challenge once the goal is met.')]
tb=doc.add_table(rows=len(legend),cols=2); tb.style='Table Grid'; tb.columns[0].width=Inches(2.0); tb.columns[1].width=Inches(4.5)
for i,(lab,g) in enumerate(legend):
    c0,c1=tb.rows[i].cells; shade(c0,'1A2332')
    p=c0.paragraphs[0]; rr=p.add_run(lab); rr.font.bold=True; rr.font.size=Pt(10); rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    p2=c1.paragraphs[0]; r2=p2.add_run(g); r2.font.size=Pt(10.5)
para('The Cornell Universal Front for each standard cites the exact Student (Lean) Deck slides so notes, vocabulary, sources, and Progress Checks stay in the deck’s sequence.',10.5,italic=True,color=GREY,before=8)
doc.add_page_break()

# ---------- Per-standard sections ----------
for code in STDS:
    md=D.STANDARDS[code]; ref=REF[code]
    doc.add_heading(f'{code} — {md["title"]}',level=1)
    # Cornell Universal Front
    doc.add_heading('Cornell Universal Front',level=2)
    ft=doc.add_table(rows=6,cols=2); ft.style='Table Grid'; ft.columns[0].width=Inches(1.7); ft.columns[1].width=Inches(4.8)
    rows=[('TN Standard',md['tn']),('I Can',md['ican']),
          ('Direct-teaching',f'Student (Lean) Deck slides {ref["dt"]} (vocabulary slide {ref["vocab"]}; primary source slide {ref["source"]}; Three Perspectives slide {ref["persp"]}; Progress Check slide {ref["progress"]}).'),
          ('Guiding questions','Who benefited?  Who bore the costs?  Who decided?  (Use these three historian questions as your note cues.)'),
          ('Key terms','; '.join(v[0] for v in D.VOCAB[code])),
          ('Primary source',f'See “Source It First,” Student Deck slide {ref["source"]}.')]
    for i,(k,v) in enumerate(rows):
        c0,c1=ft.rows[i].cells; shade(c0,'F1EFE6')
        p=c0.paragraphs[0]; rr=p.add_run(k); rr.font.bold=True; rr.font.size=Pt(10.5); rr.font.color.rgb=NAVY
        p2=c1.paragraphs[0]; r2=p2.add_run(v); r2.font.size=Pt(10.5)
    # 25/75 Cornell notes
    doc.add_heading('Cornell Notes (25 / 75)',level=2)
    ct=doc.add_table(rows=1,cols=2); ct.style='Table Grid'; ct.columns[0].width=Inches(1.9); ct.columns[1].width=Inches(4.6)
    c0,c1=ct.rows[0].cells; shade(c0,'FBF9F2')
    p=c0.paragraphs[0]; rr=p.add_run('CUES / QUESTIONS'); rr.font.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=GREY
    for cue in ['Who benefited?','Who bore the costs?','Who decided?','Key term:','Key term:']:
        cp=c0.add_paragraph(); cr=cp.add_run(cue); cr.font.size=Pt(10.5); cp.paragraph_format.space_after=Pt(14)
    p2=c1.paragraphs[0]; r2=p2.add_run('NOTES'); r2.font.bold=True; r2.font.size=Pt(9.5); r2.font.color.rgb=GREY
    for _ in range(7): c1.add_paragraph('')
    # summary
    sm=doc.add_table(rows=1,cols=1); sm.style='Table Grid'; shade(sm.rows[0].cells[0],'F1EFE6')
    sp=sm.rows[0].cells[0].paragraphs[0]; sr=sp.add_run('SUMMARY (2–3 sentences): '); sr.font.bold=True; sr.font.size=Pt(10)
    sm.rows[0].cells[0].add_paragraph(''); sm.rows[0].cells[0].add_paragraph('')
    # Key vocabulary / language support
    doc.add_heading('Key Vocabulary · Language Support',level=2)
    vt=doc.add_table(rows=1,cols=3); vt.style='Table Grid'
    hdr=vt.rows[0].cells
    for j,h in enumerate(['Term (say / Spanish)','Definition','Use it']):
        shade(hdr[j],'1A2332'); pp=hdr[j].paragraphs[0]; rr=pp.add_run(h); rr.font.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for (term,say,es,deff) in D.VOCAB[code]:
        row=vt.add_row().cells
        p=row[0].paragraphs[0]; r1=p.add_run(term+'\n'); r1.font.bold=True; r1.font.size=Pt(10.5)
        r2=p.add_run((f'say: {say}\n' if say else '')+f'ES: {es}'); r2.font.size=Pt(9); r2.font.color.rgb=GREY
        row[1].paragraphs[0].add_run(deff).font.size=Pt(10)
        row[2].paragraphs[0].add_run('').font.size=Pt(10)
    # Transfer Check (de-biased; NO answer here)
    doc.add_heading('Transfer Check',level=2)
    item=D.CFU[code]; opts,_=D.debiased_options(code)
    para(f'[{QID[code]} · DOK {item["dok"]} · pre-field-test]  {item["stem"]}',10.5,bold=True,after=4)
    for L in ['A','B','C','D']:
        pp=doc.add_paragraph(); pp.paragraph_format.left_indent=Inches(0.3); pp.paragraph_format.space_after=Pt(3)
        rr=pp.add_run(f'{L}.  '); rr.font.bold=True; rr.font.size=Pt(10.5); pp.add_run(opts[L]).font.size=Pt(10.5)
    para('RESPONSE CHOICE — answer by writing, saying/recording, or diagramming. Answer key is in the Teacher Guide (kept separate).',9.5,italic=True,color=GREY,after=4)
    doc.add_page_break()

# ---------- Representation ----------
doc.add_heading('Representation in the Industrial Era — Reached and Left Out',level=1)
para('This section complicates the industrial-era narrative: who it reached, who it excluded, and who organized in response. It is History Hack-authored instructional synthesis grounded in this unit’s sourced record — not a primary source. Verified public-domain images and named Key-Figure entries are a pending content addition.',11)
REP=[('American Indian Nations','Westward expansion and the Transcontinental Railroad (US.01) crossed treaty lands; reservation, assimilation, boarding-school, and Dawes Act policy (US.02) cut tribal holdings from about 138 million acres to 48 million by 1934.'),
     ('Chinese Laborers','Chinese workers built the most dangerous stretches of the Central Pacific (US.01) yet faced wage discrimination and, in 1882, the first federal law to bar an entire nationality — the Chinese Exclusion Act (US.07).'),
     ('African Americans','The Compromise of 1877 ended Reconstruction’s protections (US.03); Pap Singleton’s Exodusters migrated to Kansas, while Madam C.J. Walker and George Washington Carver (US.05) built enterprise and innovation under Jim Crow.'),
     ('New Immigrants & Reformers','Southern and Eastern European and Asian arrivals (US.07) met nativism and tenement poverty; reformers such as Jane Addams at Hull-House organized mutual aid across social classes.')]
for h,b in REP:
    doc.add_heading(h,level=2); para(b,11)
para('Residual note (known follow-up): sourced visuals and named Key-Figure entries for these histories remain pending.',9.5,italic=True,color=GREY,before=6)

# ---------- TOC field + settings + properties ----------
def make(tag,**a):
    e=OxmlElement(tag)
    for k,v in a.items(): e.set(qn(k),v)
    return e
def run_with(child):
    r=OxmlElement('w:r'); r.append(child); return r
for p in doc.paragraphs:
    if 'Right-click to update this table of contents' in p.text:
        pel=p._p
        for r in list(p.runs): pel.remove(r._element)
        beg=make('w:fldChar',**{'w:fldCharType':'begin'})
        instr=make('w:instrText',**{'xml:space':'preserve'}); instr.text=' TOC \\o "1-2" \\h \\z \\u '
        sep=make('w:fldChar',**{'w:fldCharType':'separate'})
        tt=OxmlElement('w:t'); tt.text='Table of contents (updates on open in Word/LibreOffice).'
        end=make('w:fldChar',**{'w:fldCharType':'end'})
        for ch in (beg,instr,sep,tt,end): pel.append(run_with(ch))
        break
sett=doc.settings.element
if sett.find(qn('w:updateFields')) is None:
    uf=OxmlElement('w:updateFields'); uf.set(qn('w:val'),'true'); sett.insert(0,uf)
cp=doc.core_properties
cp.author='TroopToTeacher Technologies LLC'; cp.last_modified_by='TroopToTeacher Technologies LLC'
cp.title='U.S. History Hack — Unit 1 Course Standard Student Workbook (Pilot)'
cp.comments='U.S. History Hack Unit 1 (Course Standard). © 2026 TroopToTeacher Technologies LLC.'
cp.created=datetime.datetime(2026,7,25); cp.modified=datetime.datetime(2026,7,25)
doc.save(OUT)
print('SAVED',OUT,'| standards:',len(STDS),'| paragraphs:',len(doc.paragraphs),'| tables:',len(doc.tables))
