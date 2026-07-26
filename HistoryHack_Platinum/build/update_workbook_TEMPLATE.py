#!/usr/bin/env python3
"""Update Cornell/anchor deck references in the Unit 5 workbook DOCX to the FINAL
49-slide Student deck, build the TOC field, add image alt text, and correct the
document properties. Reads the PRISTINE original; writes a fresh working copy."""
import re, datetime, docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
ORIG='workbook_outputs/U.S._History_Hack_Unit_5_Course_Standard_Student_Workbook_Pilot.docx'
OUT='working/Unit5_Student_Workbook_WORKING.docx'
MAP={
 'slides 1-6':'slides 2-9','slides 7-12':'slides 10-17','slides 13-18':'slides 18-25',
 'slides 19-24':'slides 26-33','slides 25-30':'slides 34-41','slides 31-36':'slides 42-49',
 'slides 2-4':'slides 4-6','slides 8-10':'slides 12-14','slides 14-16':'slides 20-22',
 'slides 20-22':'slides 28-30','slides 26-28':'slides 36-38','slides 32-34':'slides 44-46',
 'slide 6 (Three Perspectives)':'slide 8 (Three Perspectives)','slide 12 (Three Perspectives)':'slide 16 (Three Perspectives)',
 'slide 18 (Three Perspectives)':'slide 24 (Three Perspectives)','slide 24 (Three Perspectives)':'slide 32 (Three Perspectives)',
 'slide 30 (Three Perspectives)':'slide 40 (Three Perspectives)','slide 36 (Three Perspectives)':'slide 48 (Three Perspectives)',
}
keys=sorted(MAP,key=len,reverse=True); rx=re.compile('|'.join(re.escape(k) for k in keys))
def sub(t): return rx.sub(lambda m:MAP[m.group(0)],t)

import shutil; shutil.copy(ORIG,OUT)
d=docx.Document(OUT)
def walk(container):
    for p in container.paragraphs: yield p
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells: yield from walk(cell)
n=[0]
def fix_p(p):
    changed=False
    for r in p.runs:                       # run-level, single pass on each run's own text
        if rx.search(r.text):
            new=sub(r.text)
            if new!=r.text: r.text=new; changed=True
    if not changed and rx.search(p.text):  # spanning fallback ONLY when runs missed it (prevents chaining)
        if p.runs:
            p.runs[0].text=sub(p.text)
            for r in p.runs[1:]: r.text=''
            changed=True
    if changed: n[0]+=1
for p in walk(d): fix_p(p)
for sec in d.sections:
    for hf in (sec.header,sec.footer):
        for p in walk(hf): fix_p(p)

# --- image alt text (the two verified primary-source photos) ---
ALT=['Black-and-white FSA photograph by Dorothea Lange, "Migrant Mother," 1936 (Library of Congress).',
     'Black-and-white Harris & Ewing photograph of the Bonus Army camp at Anacostia Flats, Washington, D.C., 1932 (Library of Congress).']
for i,shp in enumerate(d.inline_shapes):
    try:
        shp._inline.docPr.set('descr', ALT[i] if i<len(ALT) else 'Verified public-domain photograph; see the source line beneath the image.')
    except Exception as e: print('alt warn',i,e)

# --- insert a REAL TOC field where the placeholder text sits ---
def make(tag,**a):
    e=OxmlElement(tag)
    for k,v in a.items(): e.set(qn(k),v)
    return e
def run_with(child):
    r=OxmlElement('w:r'); r.append(child); return r
toc_done=False
for p in d.paragraphs:
    if ('Right-click' in p.text and 'table of contents' in p.text.lower()) and not toc_done:
        pel=p._p
        for r in list(p.runs): pel.remove(r._element)
        beg=make('w:fldChar',**{'w:fldCharType':'begin'})
        instr=make('w:instrText',**{'xml:space':'preserve'}); instr.text=' TOC \\o "1-3" \\h \\z \\u '
        sep=make('w:fldChar',**{'w:fldCharType':'separate'})
        t=OxmlElement('w:t'); t.text='Table of contents (updates on open in Word/LibreOffice).'
        end=make('w:fldChar',**{'w:fldCharType':'end'})
        for ch in (beg,instr,sep,t,end): pel.append(run_with(ch))
        toc_done=True
print('TOC field inserted:',toc_done)
# --- also flag update-fields-on-open ---
st=d.settings.element
if st.find(qn('w:updateFields')) is None:
    uf=OxmlElement('w:updateFields'); uf.set(qn('w:val'),'true'); st.insert(0,uf)

# --- document properties (match the copyright page) ---
cp=d.core_properties
cp.author='TroopToTeacher Technologies LLC'
cp.last_modified_by='TroopToTeacher Technologies LLC'
cp.title='U.S. History Hack — Unit 5 Course Standard Student Workbook (Pilot)'
cp.comments='U.S. History Hack Unit 5 (Course Standard). © 2026 TroopToTeacher Technologies LLC.'
cp.created=datetime.datetime(2026,7,24); cp.modified=datetime.datetime(2026,7,25)

# --- append a Representation section (authored synthesis; appears in the TOC) ---
REP=[('Black Americans','The AAA’s crop-reduction payments went to landowners, pushing many Black sharecroppers and tenant farmers off the land, and Social Security (1935) at first excluded farm and domestic workers — jobs held by most Black workers. Yet FDR’s informal “Black Cabinet,” including Mary McLeod Bethune and Robert Weaver, advised the administration, and Black voters shifted toward the Democratic Party.'),
 ('Mexican Americans','Amid Depression job competition, “Mexican Repatriation” (about 1929–1936) pressured an estimated 400,000 to more than one million people of Mexican descent — many of them U.S.-born citizens — to leave for Mexico. Mexican and Filipino farmworkers in California often faced lower pay and worse conditions than white “Okie” migrants.'),
 ('Indigenous Peoples','The Indian Reorganization Act of 1934 (the “Indian New Deal”), led by John Collier, ended the Dawes-era allotment policy, restored some tribal self-government and land, and supported Native culture and religion — a major reversal, though its results were mixed and debated. A CCC Indian Division employed thousands.')]
d.add_page_break()
d.add_heading('Representation in the New Deal Era — Reached and Left Out', level=1)
d.add_paragraph('This section complicates the New Deal narrative: who it reached, who it excluded, and who organized in response. It is History Hack-authored instructional synthesis grounded in the standard historical record — not a primary source. Verified public-domain images and primary sources for these histories are a pending content addition (US.41 and US.43).')
for h,b in REP:
    d.add_heading(h, level=2); d.add_paragraph(b)
d.save(OUT)

# patch app.xml Company/Application if it names the drafting tool
import zipfile
z=zipfile.ZipFile(OUT); files={i.filename:z.read(i.filename) for i in z.infolist()}; infos=z.infolist(); z.close()
if 'docProps/app.xml' in files:
    app=files['docProps/app.xml'].decode()
    app=re.sub(r'<Company>.*?</Company>','<Company>TroopToTeacher Technologies LLC</Company>',app)
    files['docProps/app.xml']=app.encode()
    with zipfile.ZipFile(OUT,'w',zipfile.ZIP_DEFLATED) as o:
        for i in infos: o.writestr(i,files[i.filename])

print('reference-fix events:',n[0])
# verify
d2=docx.Document(OUT); alltext='\n'.join(p.text for p in walk(d2))
print('OLD refs remaining:',[k for k in MAP if k in alltext] or 'NONE')
for p in d2.paragraphs:
    if 'Lean Student Deck slides' in p.text: print('  opener:',p.text.strip()[:100])
print('creator:',d2.core_properties.author,'| created:',d2.core_properties.created)
